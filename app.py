from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
import tempfile
import zipfile
import io
from pathlib import Path
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = "renumerador-recibos"

ALLOWED_EXTENSIONS = {".docx", ".zip"}


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def encontrar_ultimo_numero(doc: Document) -> str:
    numeros = []

    for i in range(len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "NOTA DE BALCÃO":
            if i + 1 < len(doc.paragraphs):
                num = doc.paragraphs[i + 1].text.strip()
                if num.isdigit():
                    numeros.append(num)

    if numeros:
        return max(numeros, key=lambda x: int(x))
    return "0"


def renumerar_documento(caminho_entrada: str, caminho_saida: str) -> tuple[int, str]:
    doc = Document(caminho_entrada)

    ultimo = encontrar_ultimo_numero(doc)
    tamanho = len(ultimo)
    numero_atual = int(ultimo) + 1

    i = 0
    alterados = 0

    while i < len(doc.paragraphs):
        if doc.paragraphs[i].text.strip() == "NOTA DE BALCÃO":
            if i + 1 < len(doc.paragraphs):
                texto_num = doc.paragraphs[i + 1].text.strip()
                if texto_num.isdigit():
                    novo_num = str(numero_atual).zfill(tamanho)
                    doc.paragraphs[i + 1].text = novo_num
                    numero_atual += 1
                    alterados += 1
                    i += 1
        i += 1

    Path(caminho_saida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(caminho_saida)
    return alterados, ultimo


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        arquivo = request.files.get("arquivo")

        if not arquivo or not arquivo.filename:
            flash("Selecione um arquivo .docx ou .zip.")
            return redirect(url_for("index"))

        if not allowed_file(arquivo.filename):
            flash("Formato inválido. Envie um arquivo .docx ou .zip.")
            return redirect(url_for("index"))

        with tempfile.TemporaryDirectory() as temp_dir_str:
            temp_dir = Path(temp_dir_str)
            entrada_dir = temp_dir / "entrada"
            saida_dir = temp_dir / "saida"
            entrada_dir.mkdir(parents=True, exist_ok=True)
            saida_dir.mkdir(parents=True, exist_ok=True)

            nome_seguro = secure_filename(arquivo.filename) or "upload"
            caminho_upload = temp_dir / nome_seguro
            arquivo.save(caminho_upload)

            if caminho_upload.suffix.lower() == ".zip":
                try:
                    with zipfile.ZipFile(caminho_upload, "r") as zip_ref:
                        zip_ref.extractall(entrada_dir)
                except zipfile.BadZipFile:
                    flash("O arquivo ZIP enviado está corrompido ou inválido.")
                    return redirect(url_for("index"))
            else:
                destino = entrada_dir / nome_seguro
                destino.write_bytes(caminho_upload.read_bytes())

            arquivos_docx = [
                p for p in entrada_dir.rglob("*.docx")
                if not p.name.startswith("~$")
            ]

            if not arquivos_docx:
                flash("Nenhum arquivo .docx foi encontrado para processar.")
                return redirect(url_for("index"))

            total_arquivos = 0
            total_recibos = 0
            relatorio = []

            for caminho in arquivos_docx:
                relativo = caminho.relative_to(entrada_dir)
                destino = saida_dir / relativo

                try:
                    alterados, ultimo = renumerar_documento(str(caminho), str(destino))
                    total_arquivos += 1
                    total_recibos += alterados
                    relatorio.append(
                        f"{relativo.as_posix()} | último encontrado: {ultimo} | recibos renumerados: {alterados}"
                    )
                except Exception as e:
                    relatorio.append(
                        f"{relativo.as_posix()} | erro: {str(e)}"
                    )

            relatorio_path = saida_dir / "relatorio_processamento.txt"
            relatorio_path.write_text(
                "RENumerador de Recibos - Relatório de Processamento\n\n"
                f"Arquivos processados: {total_arquivos}\n"
                f"Recibos renumerados: {total_recibos}\n\n"
                + "\n".join(relatorio),
                encoding="utf-8"
            )

            if len(arquivos_docx) == 1:
                unico_saida = next((p for p in saida_dir.rglob("*.docx")), None)
                if unico_saida:
                    buffer = io.BytesIO(unico_saida.read_bytes())
                    buffer.seek(0)
                    return send_file(
                        buffer,
                        as_attachment=True,
                        download_name=unico_saida.name,
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            zip_path = temp_dir / "recibos_renumerados.zip"
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zip_out:
                for arquivo_saida in saida_dir.rglob("*"):
                    if arquivo_saida.is_file():
                        zip_out.write(arquivo_saida, arquivo_saida.relative_to(saida_dir))

            zip_buffer = io.BytesIO(zip_path.read_bytes())
            zip_buffer.seek(0)

            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name="recibos_renumerados.zip",
                mimetype="application/zip"
            )

    return render_template("index.html")


if __name__ == "__main__":
    app.run(debug=True)
